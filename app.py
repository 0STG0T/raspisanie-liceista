from flask import Flask, render_template, request, redirect, url_for, session
from docx import Document
import os
from werkzeug.utils import secure_filename
from datetime import datetime
import locale
from datetime import datetime
from babel.dates import format_date
from datetime import datetime
import telegram

UPLOAD_FOLDER = './uploaded/'  # Задайте путь к директории для загрузки файлов
ALLOWED_EXTENSIONS = {'docx'}

classes = ['11А', '11Б', '10А', '10Б', '9А', '9Б', '8А', '8Б', '7А', '7Б']  # Добавьте другие классы

chat_ids = ['803006666']
TOKEN = '6361532905:AAFGIBmPOBrgxk77F-mBJGlQz-7QAZaxJKk'

# Установка локали для корректного отображения дней недели на русском языке
locale.setlocale(locale.LC_TIME, 'ru_RU.UTF-8')

def send_message_to_all_users(text):
    bot = telegram.Bot(token=TOKEN)
    for chat_id in chat_ids:
        bot.send_message(chat_id=chat_id, text=text)

def get_schedule_for_day(day, file_path):
    doc = Document(file_path)
    schedule = {}
    day = day.lower()

    # Список для хранения названий классов
    classes = []

    # Флаг, показывающий, найден ли нужный день в таблице
    day_found = False

    for table in doc.tables:
        # Если классы еще не сохранены, сохраняем их из строки с понедельником
        if not classes:
            for cell in table.rows[0].cells[1:]:
                class_name = cell.text.strip()
                if class_name:  # Убедимся, что ячейка не пустая
                    classes.append(class_name)

        # Обрабатываем строки таблицы
        for row in table.rows:
            first_cell_text = row.cells[0].text.strip().lower()
            if first_cell_text in ['пн', 'вт', 'ср', 'чт', 'пт', 'сб', 'вс']:
                if first_cell_text == day:
                    day_found = True
                elif day_found:
                    break  # Если следующий день найден, прерываем обработку

            if day_found:
                time = row.cells[0].text.strip()
                for class_index, class_name in enumerate(classes):
                    subject_index = class_index + 1  # Смещаем индекс на 1, так как первая ячейка - время
                    if subject_index < len(row.cells):
                        subject = row.cells[subject_index].text.strip()
                        if subject:  # Если есть предмет, добавляем его в расписание
                            if class_name not in schedule:
                                schedule[class_name] = []
                            schedule[class_name].append((time, subject))
    return schedule

def get_schedule_for_week(class_name, file_path):
    doc = Document(file_path)
    weekly_schedule = {}

    # Список для хранения названий классов, будет заполнен из первой строки таблицы
    classes = None
    current_day = None  # Текущий день недели

    for table in doc.tables:
        if classes is None:
            classes = [cell.text.strip() for cell in table.rows[0].cells[1:]]

        class_index = None
        if class_name in classes:
            class_index = classes.index(class_name) + 1  # +1 так как первая ячейка - время

        for row in table.rows:
            first_cell_text = row.cells[0].text.strip().lower()
            if first_cell_text in ['пн', 'вт', 'ср', 'чт', 'пт', 'сб', 'вс']:
                current_day = first_cell_text

            if current_day and class_index:
                time = row.cells[0].text.strip()
                subject = row.cells[class_index].text.strip() if class_index < len(row.cells) else ''
                if subject:
                    if current_day not in weekly_schedule:
                        weekly_schedule[current_day] = []
                    weekly_schedule[current_day].append((time, subject))

    return weekly_schedule


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

weekdays = {
    0: 'Понедельник',
    1: 'Вторник',
    2: 'Среда',
    3: 'Четверг',
    4: 'Пятница',
    5: 'Суббота',
    6: 'Воскресенье'
}

app = Flask(__name__)
app.secret_key = '42'  # Замените на ваш собственный секретный ключ
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Список классов и их букв


@app.before_request
def setup():
    # Если пользователь не аутентифицирован и не находится на странице входа или загрузки расписания, перенаправляем на страницу входа
    if 'user_id' not in session and request.endpoint not in ['login', 'upload_schedule']:
        return redirect(url_for('login'))

@app.route('/')
def home():
    # Получаем текущий день недели
    current_day = format_date(datetime.now(), format='EEEE').lower()

    # Переводим день недели на русский
    days_mapping = {'monday': 'пн', 'tuesday': 'вт', 'wednesday': 'ср', 'thursday': 'чт', 'friday': 'пт', 'saturday': 'сб', 'sunday': 'вс'}
    current_day_short = days_mapping[current_day]
    # Предположим, что class_name хранится в сессии или определяется иным способом
    class_name = session.get('user_id')  # Пример, замените на реальное значение
    class_name = class_name[:-1] + ' ' + class_name[-1] 
    # Получаем расписание на текущий день для данного класса
    schedule_data = get_schedule_for_day(current_day_short, r'.\uploaded\raspur.docx')[class_name]
    
    return render_template('front.html', schedule=schedule_data, day=current_day.capitalize(), class_name=class_name)
    

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        user_id = request.form['user_id']
        if user_id in classes:
            session['user_id'] = user_id  # Устанавливаем сессию для пользователя
            return redirect(url_for('home'))
        elif user_id == 'teacher':
            return redirect(url_for('upload_schedule'))
    return render_template('login.html')

@app.route('/upload_schedule', methods=['GET', 'POST'])
def upload_schedule():
    if request.method == 'POST':
        file = request.files['file']
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(save_path)
            app.config['last_uploaded_file'] = save_path
            return redirect(url_for('home'))
    return render_template('upload_schedule.html')

@app.route('/schedule')
def schedule():
    # Путь к файлу расписания
    file_path = r'.\uploaded\raspur.docx'
    class_name = session.get('user_id')  # Например, '11 А'
    class_name = class_name[:-1] + ' ' + class_name[-1] 
    weekly_schedule = get_schedule_for_week(class_name, file_path)

    return render_template('schedule.html', weekly_schedule=weekly_schedule, class_name=class_name)

@app.route('/video_lessons')
def video_lessons():
    # Список URL видео с YouTube
    videos = [
        'https://www.youtube.com/embed/N0ppXbMmSPA',
        # Добавьте больше видео здесь
    ]
    return render_template('video_lessons.html', videos=videos)

@app.route('/feedback')
def feedback():
    return render_template('feedback.html')

@app.route('/logout')
def logout():
    # Удаляем сессию пользователя
    session.pop('user_id', None)
    return redirect(url_for('login'))

if __name__ == '__main__':
    app.run(debug=True)
